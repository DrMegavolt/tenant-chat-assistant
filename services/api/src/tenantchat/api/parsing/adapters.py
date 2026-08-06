"""Format adapters: bytes in a media type, structured text and locations out.

Each adapter is one explicit implementation for one family of documents, with
its own ``version`` identifier so the generation record (`OBS-004`) can name
exactly which parser produced a document's chunks. Adapters own their
validation: corrupt, encrypted, and empty documents raise
:class:`~tenantchat.core.errors.ValidationError` here, and the ingestion
handler turns that into the permanent scan-rejected job failure.

Markdown, HTML, and plain text are parsed with the standard library and a
deliberately small hand-rolled parser — the pipeline needs headings, code
fences, and table rows, not full CommonMark — while PDF and DOCX use
``pypdf`` and ``python-docx``, the smallest maintained pure-Python parsers
for those formats.
"""

from __future__ import annotations

import io
import re
import zipfile
from collections.abc import Iterator, Sequence
from html.parser import HTMLParser
from typing import Protocol, cast

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from pypdf.generic import Destination

from tenantchat.api.parsing.locations import ChunkLocation, ParsedDocument, SourceBlock
from tenantchat.api.parsing.scan import scan_bytes, scan_text
from tenantchat.core.errors import ValidationError

_ATX_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_SETEXT_UNDERLINE = re.compile(r"^(={3,}|-{3,})$")
_FENCE_MARKER = ("```", "~~~")
_DELIMITER_CELL_RE = re.compile(r":?-{1,}:?")

_MARKDOWN_MEDIA_TYPES = frozenset({"text/markdown", "text/x-markdown"})
_HTML_MEDIA_TYPES = frozenset({"text/html"})
_TEXT_MEDIA_TYPES = frozenset({"text/plain"})
_PDF_MEDIA_TYPES = frozenset({"application/pdf"})
_DOCX_MEDIA_TYPES = frozenset(
    {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
)


class ParserAdapter(Protocol):
    """One versioned format adapter: bytes in, structured blocks out."""

    version: str
    media_types: frozenset[str]

    def parse(self, content: bytes, *, title: str, media_type: str) -> ParsedDocument: ...


def _push_heading(headings: list[str], depth: int, text: str) -> list[str]:
    """Update the heading stack: deeper pushes, same level replaces, shallower pops.

    A level jump (h1 followed by h3) makes the deeper heading a child of the
    existing stack, exactly as a document reader would expect.
    """
    while len(headings) >= depth:
        headings.pop()
    headings.append(text)
    return headings


def _is_table_delimiter_row(line: str) -> bool:
    """Whether a pipe row is a markdown table's header separator.

    A delimiter row's cells contain only dashes, colons, and pipes; dropping
    it keeps the table's text clean without losing the rows that carry
    content.
    """
    body = line.strip().strip("|").strip()
    if not body:
        return False
    return all(_DELIMITER_CELL_RE.fullmatch(cell.strip()) for cell in body.split("|"))


def _clean_lines(text: str) -> list[str]:
    """Normalize extracted text to stripped, non-blank lines.

    Extractors differ in where they put newlines; this is what makes their
    output comparable and deterministic across library versions.
    """
    return [line for line in (line.strip() for line in text.splitlines()) if line]


class MarkdownParser:
    """The subset of CommonMark the pipeline needs: ATX and setext headings,
    fenced code blocks (whose contents are never headings), and pipe tables.
    Everything else is body text.
    """

    version = "markdown.v1"
    media_types = _MARKDOWN_MEDIA_TYPES

    def parse(self, content: bytes, *, title: str, media_type: str) -> ParsedDocument:
        scan_text(content)
        blocks: list[SourceBlock] = []
        headings: list[str] = []
        buffer: list[str] = []
        fence: str | None = None
        previous_line: str | None = None

        def flush() -> None:
            if buffer:
                blocks.append(
                    SourceBlock(
                        location=ChunkLocation(section_path=tuple(headings)),
                        text="\n".join(buffer),
                    )
                )
                buffer.clear()

        for line in content.decode("utf-8").splitlines():
            stripped = line.strip()
            if fence is not None:
                if stripped.startswith(fence):
                    fence = None
                else:
                    buffer.append(line)
                previous_line = None
                continue
            if stripped.startswith(_FENCE_MARKER):
                fence = stripped[:3]
                previous_line = None
                continue

            match = _ATX_HEADING_RE.match(stripped)
            if match:
                heading = match.group(2).strip().rstrip("#").strip()
                if heading:
                    flush()
                    headings = _push_heading(headings, len(match.group(1)), heading)
                    buffer.append(heading)
                previous_line = None
                continue
            if _SETEXT_UNDERLINE.match(stripped) and previous_line is not None:
                level = 1 if stripped.startswith("=") else 2
                # The underlined line is the heading itself: take it out of
                # the body buffer so it lands under its own heading path,
                # exactly as an ATX heading does.
                if buffer and buffer[-1] == previous_line:
                    buffer.pop()
                flush()
                headings = _push_heading(headings, level, previous_line)
                buffer.append(previous_line)
                previous_line = None
                continue
            if _is_table_delimiter_row(line):
                previous_line = None
                continue
            if stripped:
                buffer.append(stripped)
                previous_line = stripped
            else:
                previous_line = None
        flush()

        if not blocks:
            raise ValidationError(detail="document has no text content")
        return ParsedDocument(
            title=title, media_type=media_type, parser_version=self.version, blocks=tuple(blocks)
        )


_BLOCK_END_TAGS = frozenset(
    {"p", "div", "section", "article", "li", "ul", "ol", "br", "tr", "blockquote", "pre"}
)
_HEADING_LEVELS = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}
_IGNORED_SECTIONS = frozenset({"head", "script", "style"})


class _HtmlExtractor(HTMLParser):
    """Flows HTML text into lines, capturing heading hierarchy and ids.

    Block-level elements start a new line; inline elements flow into the
    current line, so ``<p>Hello <b>world</b></p>`` yields one line instead of
    two. A table row becomes one line with its cells joined by pipes, so a
    row's rendering does not depend on whether a heading text shares the
    line.
    """

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[SourceBlock] = []
        self._headings: list[str] = []
        self._lines: list[str] = []
        self._line: list[str] = []
        self._ignored_depth = 0
        self._heading_level: int | None = None
        self._heading_words: list[str] = []
        self._heading_anchor: str | None = None
        # The anchor of the heading that opened the block being flushed;
        # every block under that heading inherits it.
        self._block_anchor: str | None = None
        self._row_cells: list[str] = []
        self._current_cell: list[str] | None = None

    def _end_line(self) -> None:
        if self._line:
            self._lines.append(" ".join(self._line))
            self._line = []

    def _flush(self) -> None:
        self._end_line()
        if self._lines:
            self.blocks.append(
                SourceBlock(
                    location=ChunkLocation(
                        section_path=tuple(self._headings), anchor=self._block_anchor
                    ),
                    text="\n".join(self._lines),
                )
            )
            self._lines = []

    def finish(self) -> None:
        """Flush the trailing block; call once after the whole document fed.

        ``feed`` only flushes at a heading boundary, so the last section of
        the document would otherwise be dropped for want of a following
        heading.
        """
        self._flush()

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in _IGNORED_SECTIONS:
            self._ignored_depth += 1
            return
        if self._ignored_depth:
            return
        if tag in _HEADING_LEVELS:
            self._flush()
            self._heading_level = _HEADING_LEVELS[tag]
            self._heading_words = []
            self._heading_anchor = dict(attrs).get("id")
        elif tag == "br":
            self._end_line()
        elif tag in {"td", "th"}:
            self._current_cell = []

    def handle_endtag(self, tag: str) -> None:
        if tag in _IGNORED_SECTIONS:
            if self._ignored_depth:
                self._ignored_depth -= 1
            return
        if self._ignored_depth:
            return
        level = _HEADING_LEVELS.get(tag)
        if level is not None:
            heading = " ".join(self._heading_words)
            anchor = self._heading_anchor
            self._heading_level = None
            self._heading_words = []
            self._heading_anchor = None
            if heading:
                self._headings = _push_heading(self._headings, level, heading)
                self._block_anchor = anchor
                self._line.append(heading)
        elif tag in {"td", "th"}:
            if self._current_cell is not None:
                self._row_cells.append(" ".join(self._current_cell))
                self._current_cell = None
        elif tag == "tr":
            self._end_line()
            if self._row_cells:
                self._line.append(" | ".join(self._row_cells))
                self._row_cells = []
            self._end_line()
        elif tag in _BLOCK_END_TAGS:
            self._end_line()

    def handle_data(self, data: str) -> None:
        if self._ignored_depth:
            return
        words = data.split()
        if not words:
            return
        if self._heading_level is not None:
            self._heading_words.extend(words)
        elif self._current_cell is not None:
            self._current_cell.extend(words)
        else:
            self._line.extend(words)


class HtmlParser:
    """Parses HTML into heading-ordered text blocks with heading-id anchors."""

    version = "html.v1"
    media_types = _HTML_MEDIA_TYPES

    def parse(self, content: bytes, *, title: str, media_type: str) -> ParsedDocument:
        scan_text(content)
        extractor = _HtmlExtractor()
        extractor.feed(content.decode("utf-8"))
        extractor.finish()
        if not extractor.blocks:
            raise ValidationError(detail="document has no text content")
        return ParsedDocument(
            title=title,
            media_type=media_type,
            parser_version=self.version,
            blocks=tuple(extractor.blocks),
        )


class TextParser:
    """Plain text: one document-level block with no structure."""

    version = "text.v1"
    media_types = _TEXT_MEDIA_TYPES

    def parse(self, content: bytes, *, title: str, media_type: str) -> ParsedDocument:
        scan_text(content)
        lines = _clean_lines(content.decode("utf-8"))
        if not lines:
            raise ValidationError(detail="document has no text content")
        return ParsedDocument(
            title=title,
            media_type=media_type,
            parser_version=self.version,
            blocks=(SourceBlock(location=ChunkLocation(), text="\n".join(lines)),),
        )


class PdfParser:
    """Parses PDFs page by page, threading outline headings onto their pages.

    A document outline item names the section its page starts, and applies to
    every page until the next item at the same or shallower depth; documents
    without a usable outline fall back to page-number locations. Blank pages
    produce no blocks but keep their absolute page numbers.
    """

    version = "pdf.v1"
    media_types = _PDF_MEDIA_TYPES

    def parse(self, content: bytes, *, title: str, media_type: str) -> ParsedDocument:
        scan_bytes(content)
        try:
            reader = PdfReader(io.BytesIO(content))
        except PdfReadError as exc:
            raise ValidationError(detail="document is corrupt") from exc
        if reader.is_encrypted:
            raise ValidationError(detail="document is encrypted")

        headings_by_page = self._outline_headings(reader)
        blocks: list[SourceBlock] = []
        for page_number, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception as exc:
                raise ValidationError(detail="document is corrupt") from exc
            lines = _clean_lines(page_text)
            if not lines:
                continue
            blocks.append(
                SourceBlock(
                    location=ChunkLocation(
                        section_path=headings_by_page.get(page_number, ()), page=page_number
                    ),
                    text="\n".join(lines),
                )
            )
        if not blocks:
            raise ValidationError(detail="document has no text content")
        return ParsedDocument(
            title=title, media_type=media_type, parser_version=self.version, blocks=tuple(blocks)
        )

    @staticmethod
    def _outline_headings(reader: PdfReader) -> dict[int, tuple[str, ...]]:
        """Map page numbers to heading paths, threaded from the outline.

        Returns an empty mapping when the outline is absent or cannot be
        resolved, so a PDF without bookmarks still parses with page-level
        locations.
        """
        if not reader.outline:
            return {}
        entries: list[tuple[int, int, str]] = []
        for item, depth in PdfParser._walk_outline(reader.outline):
            title = getattr(item, "title", None)
            if not isinstance(title, str) or not title.strip():
                continue
            page = PdfParser._outline_page(reader, item)
            if page is None:
                continue
            entries.append((page, depth, title.strip()))
        entries.sort(key=lambda entry: (entry[0], entry[1]))

        headings: dict[int, tuple[str, ...]] = {}
        stack: list[tuple[int, str]] = []
        index = 0
        for page in range(1, len(reader.pages) + 1):
            while index < len(entries) and entries[index][0] == page:
                _, depth, title = entries[index]
                while stack and stack[-1][0] >= depth:
                    stack.pop()
                if title:
                    stack.append((depth, title))
                index += 1
            headings[page] = tuple(title for _, title in stack)
        return headings

    @staticmethod
    def _outline_page(reader: PdfReader, item: object) -> int | None:
        """The 1-based page an outline item points at, or ``None``.

        pypdf numbers destinations from zero, so the returned page is
        incremented to match the 1-based locations every other block carries.
        Outline destinations are unreliable across producers — some name a
        page number, some an indirect object that fails to resolve — so an
        unresolvable item is skipped rather than failing the whole document.
        """
        if not hasattr(item, "title"):
            return None
        try:
            page = reader.get_destination_page_number(cast(Destination, item))
        except Exception:
            return None
        if page is None:
            return None
        return page + 1

    @staticmethod
    def _walk_outline(outline: Sequence[object], depth: int = 0) -> Iterator[tuple[object, int]]:
        """Yield ``(item, depth)`` pairs from pypdf's nested outline structure.

        pypdf types every outline entry as a ``Destination`` (itself a dict)
        and nests sub-outlines as lists, so the walk recurses into lists and
        yields every non-list entry at its current depth.
        """
        for entry in outline:
            if isinstance(entry, list):
                yield from PdfParser._walk_outline(entry, depth + 1)
            else:
                yield entry, depth


class DocxParser:
    """Parses DOCX in document order: headings, body paragraphs, and tables.

    Heading levels come from the built-in style id (``Heading1`` …), which is
    locale-independent, so a French or Spanish document yields the same
    heading hierarchy as an English one.
    """

    version = "docx.v1"
    media_types = _DOCX_MEDIA_TYPES

    def parse(self, content: bytes, *, title: str, media_type: str) -> ParsedDocument:
        scan_bytes(content)
        try:
            document = Document(io.BytesIO(content))
        except (zipfile.BadZipFile, KeyError, ValueError, PackageNotFoundError) as exc:
            raise ValidationError(detail="document is corrupt") from exc

        blocks: list[SourceBlock] = []
        headings: list[str] = []
        buffer: list[str] = []

        def flush() -> None:
            if buffer:
                blocks.append(
                    SourceBlock(
                        location=ChunkLocation(section_path=tuple(headings)),
                        text="\n".join(buffer),
                    )
                )
                buffer.clear()

        body = document.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                paragraph = Paragraph(child, document)
                text = paragraph.text.strip()
                level = DocxParser._heading_level(paragraph)
                if level is not None and text:
                    flush()
                    headings = _push_heading(headings, level, text)
                    buffer.append(text)
                elif text:
                    buffer.append(text)
            elif child.tag == qn("w:tbl"):
                for row in Table(child, document).rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        buffer.append(" | ".join(cells))
        flush()

        if not blocks:
            raise ValidationError(detail="document has no text content")
        return ParsedDocument(
            title=title, media_type=media_type, parser_version=self.version, blocks=tuple(blocks)
        )

    @staticmethod
    def _heading_level(paragraph: Paragraph) -> int | None:
        """The heading level of a paragraph, or ``None`` for body text."""
        style = paragraph.style
        style_id = getattr(style, "style_id", None)
        if not isinstance(style_id, str) or not style_id.startswith("Heading"):
            return None
        suffix = style_id[len("Heading") :]
        if not suffix.isdigit():
            return None
        return int(suffix)
