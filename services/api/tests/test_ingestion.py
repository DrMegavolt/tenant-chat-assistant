"""The durable ingestion job: parse, scan, chunk, embed, index, record.

Hermetic by construction: every external dependency is the in-memory fake
(``InMemoryKnowledgeStore``, ``InMemoryIndexIntegrityStore``,
``MemoryObjectStore``, ``InMemorySearchIndex``, ``ScriptedEmbedder``).
"""

from __future__ import annotations

import asyncio
import uuid
from collections.abc import Sequence
from datetime import UTC, datetime, timedelta

import pytest

from tenantchat.api.index_integrity import InMemoryIndexIntegrityStore
from tenantchat.api.ingestion import (
    EMBED_BATCH_SIZE,
    IngestionDependencies,
    generation_id_for,
    ingestion_handler,
    submit_ingestion,
)
from tenantchat.api.jobs import InMemoryJobStore, JobExecutionError
from tenantchat.api.parsing import (
    CHUNK_OVERLAP,
    CHUNK_TOKENS,
    chunk_document,
    chunk_text,
    parse_document,
)
from tenantchat.api.search import (
    EmbeddingResult,
    EmbeddingUnavailableError,
    IndexedChunk,
    InMemorySearchIndex,
    ScriptedEmbedder,
    SearchIndexOperationError,
)
from tenantchat.api.storage import MemoryObjectStore, StorageKey
from tenantchat.api.store import InMemoryKnowledgeStore
from tenantchat.core.errors import ValidationError
from tenantchat.core.indexing import GenerationStatus
from tenantchat.core.knowledge import (
    ContentChecksum,
    KnowledgeDomain,
    RetrievalAudience,
    RetrievalContext,
    SourceKind,
)
from tenantchat.core.lifecycle import IndexingState, VersionState

FINANCING = KnowledgeDomain.parse("financing")
NOW = datetime(2026, 8, 5, 12, 0, tzinfo=UTC)
# Long enough to span two 650-token chunks, which exercises chunk boundaries
# and the chunk-count verification path.
CONTENT = (
    b"# Financing options\n\n"
    + (" ".join(f"term-{index}" for index in range(800)).encode())
    + b"\n"
)

TWO_CHUNKS = len(
    chunk_document(
        parse_document(CONTENT, media_type="text/markdown", title="Financing plan terms"),
        chunk_tokens=CHUNK_TOKENS,
        overlap_tokens=CHUNK_OVERLAP,
    )
)


class Pipeline:
    """One wired ingestion pipeline over hermetic fakes."""

    def __init__(self) -> None:
        self.knowledge = InMemoryKnowledgeStore()
        self.generations = InMemoryIndexIntegrityStore()
        self.storage = MemoryObjectStore()
        self.index = InMemorySearchIndex()
        self.embedder = ScriptedEmbedder()
        self.jobs = InMemoryJobStore()
        self.dependencies = IngestionDependencies(
            knowledge=self.knowledge,
            generations=self.generations,
            storage=self.storage,
            index=self.index,
            embedder=self.embedder,
        )

    async def register_source(self, *, tenant_id: str = "clearview") -> uuid.UUID:
        source = await self.knowledge.register_source(
            tenant_id,
            domain=FINANCING,
            kind=SourceKind.UPLOAD,
            display_name="Financing brochures",
        )
        return source.source_id

    async def upload_and_stage(
        self,
        content: bytes = CONTENT,
        *,
        tenant_id: str = "clearview",
        external_key: str = "financing-options.md",
        media_type: str = "text/markdown",
    ) -> tuple[uuid.UUID, uuid.UUID]:
        """Store bytes and stage a draft, exactly as the upload route does."""
        source_id = await self.register_source(tenant_id=tenant_id)
        checksum = ContentChecksum.of(content)
        key = StorageKey.build(
            tenant_id=tenant_id,
            source_id=source_id,
            external_key=external_key,
            checksum=checksum.value,
        )
        await self.storage.put(key, content)
        document = await self.knowledge.stage_version(
            tenant_id,
            source_id=source_id,
            external_key=external_key,
            title="Financing plan terms",
            checksum=checksum,
            byte_size=len(content),
            media_type=media_type,
            storage_key=str(key),
        )
        version = document.version_with_checksum(checksum)
        assert version is not None
        return version.document_id, version.version_id

    async def make_indexable(
        self, version_id: uuid.UUID, *, tenant_id: str = "clearview", at: datetime = NOW
    ) -> None:
        """Approve and publish, which is what makes a version indexable."""
        await self.knowledge.approve(tenant_id, version_id, approved_by="ops@example", at=at)
        await self.knowledge.publish(tenant_id, version_id, at=at)

    async def run_job(self, tenant_id: str, version_id: uuid.UUID) -> None:
        job = await submit_ingestion(self.jobs, tenant_id=tenant_id, version_id=version_id)
        await ingestion_handler(self.dependencies)(job)

    async def retrievable(self, *, tenant_id: str = "clearview") -> bool:
        context = RetrievalContext(
            tenant_id=tenant_id,
            domain=FINANCING,
            audience=RetrievalAudience.VISITOR,
            moment=NOW + timedelta(seconds=1),
        )
        return bool(await self.knowledge.retrievable_versions(context))


def test_a_published_version_is_ingested_end_to_end() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        _, version_id = await pipeline.upload_and_stage()
        await pipeline.make_indexable(version_id)

        await pipeline.run_job("clearview", version_id)

        document = await pipeline.knowledge.document_for_version("clearview", version_id)
        version = document.version(version_id)
        assert version.indexing_state is IndexingState.INDEXED
        assert version.state is VersionState.PUBLISHED
        generation = await pipeline.generations.generation("clearview", version_id)
        assert generation is not None
        assert generation.status is GenerationStatus.COMPLETE
        assert generation.chunk_count == generation.indexed_chunk_count == TWO_CHUNKS
        assert generation.parser_version == "markdown.v1"
        assert generation.chunker_version == "token-window.v2"
        assert generation.embedding_model == "scripted-embedder.v1"
        assert (
            await pipeline.index.active_chunk_count(tenant_id="clearview", version_id=version_id)
            == TWO_CHUNKS
        )
        assert await pipeline.retrievable()

    asyncio.run(scenario())


def test_the_generation_id_is_deterministic_per_tenant_and_version() -> None:
    version_id = uuid.uuid4()
    assert generation_id_for("clearview", version_id) == generation_id_for("clearview", version_id)
    assert generation_id_for("clearview", version_id) != generation_id_for("apex", version_id)


def test_reingesting_unchanged_content_is_an_idempotent_no_op() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        _, version_id = await pipeline.upload_and_stage()
        await pipeline.make_indexable(version_id)
        await pipeline.run_job("clearview", version_id)
        chunk_ids_before = set(pipeline.index._chunks)

        await pipeline.run_job("clearview", version_id)

        assert set(pipeline.index._chunks) == chunk_ids_before
        assert (
            await pipeline.index.active_chunk_count(tenant_id="clearview", version_id=version_id)
            == TWO_CHUNKS
        )

    asyncio.run(scenario())


def test_a_duplicate_upload_creates_no_new_revision_and_no_new_job() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        _, first_version = await pipeline.upload_and_stage()
        _, second_version = await pipeline.upload_and_stage()
        assert second_version == first_version

        first_job = await submit_ingestion(
            pipeline.jobs, tenant_id="clearview", version_id=first_version
        )
        second_job = await submit_ingestion(
            pipeline.jobs, tenant_id="clearview", version_id=second_version
        )
        assert first_job.job_id == second_job.job_id

    asyncio.run(scenario())


def test_a_draft_version_is_retryable_until_it_is_approved() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        _, version_id = await pipeline.upload_and_stage()

        with pytest.raises(JobExecutionError) as captured:
            await pipeline.run_job("clearview", version_id)
        assert captured.value.error_code == "ingestion_version_not_indexable"
        assert captured.value.retryable

        await pipeline.make_indexable(version_id)
        await pipeline.run_job("clearview", version_id)
        document = await pipeline.knowledge.document_for_version("clearview", version_id)
        assert document.version(version_id).indexing_state is IndexingState.INDEXED

    asyncio.run(scenario())


def test_an_unknown_or_cross_tenant_version_is_a_permanent_missing_failure() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        _, version_id = await pipeline.upload_and_stage()
        await pipeline.make_indexable(version_id)

        with pytest.raises(JobExecutionError) as missing:
            await pipeline.run_job("clearview", uuid.uuid4())
        assert missing.value.error_code == "ingestion_version_missing"
        assert not missing.value.retryable

        with pytest.raises(JobExecutionError) as cross_tenant:
            await pipeline.run_job("apex", version_id)
        assert cross_tenant.value.error_code == "ingestion_version_missing"
        assert not cross_tenant.value.retryable

    asyncio.run(scenario())


def test_missing_stored_content_fails_retryably_and_leaves_a_failed_generation() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        _, version_id = await pipeline.upload_and_stage()
        await pipeline.make_indexable(version_id)
        version = (await pipeline.knowledge.document_for_version("clearview", version_id)).version(
            version_id
        )
        await pipeline.storage.delete(StorageKey.parse(version.storage_key))

        with pytest.raises(JobExecutionError) as captured:
            await pipeline.run_job("clearview", version_id)
        assert captured.value.error_code == "ingestion_content_missing"
        assert captured.value.retryable

        document = await pipeline.knowledge.document_for_version("clearview", version_id)
        assert document.version(version_id).indexing_state is IndexingState.FAILED
        generation = await pipeline.generations.generation("clearview", version_id)
        assert generation is not None and generation.status is GenerationStatus.FAILED
        assert not await pipeline.retrievable()

    asyncio.run(scenario())


def test_a_mid_index_failure_cleans_up_partial_chunks_on_retry() -> None:
    """The acceptance criterion: retry without duplicate active chunks."""

    async def scenario() -> None:
        pipeline = Pipeline()
        _, version_id = await pipeline.upload_and_stage()
        await pipeline.make_indexable(version_id)
        original = pipeline.embedder

        class FailingEmbedder(ScriptedEmbedder):
            attempts = 0

            async def embed(self, texts: Sequence[str]) -> EmbeddingResult:
                self.attempts += 1
                if self.attempts == 1:
                    raise EmbeddingUnavailableError("provider timeout")
                return await super().embed(texts)

        pipeline.dependencies.embedder = FailingEmbedder()

        with pytest.raises(JobExecutionError) as captured:
            await pipeline.run_job("clearview", version_id)
        assert captured.value.error_code == "ingestion_embedding_unavailable"
        assert captured.value.retryable

        pipeline.dependencies.embedder = original
        await pipeline.run_job("clearview", version_id)

        assert (
            await pipeline.index.active_chunk_count(tenant_id="clearview", version_id=version_id)
            == TWO_CHUNKS
        )
        assert (
            len(pipeline.index._chunks) == TWO_CHUNKS
        ), "a retry must not leave duplicate active chunks"
        document = await pipeline.knowledge.document_for_version("clearview", version_id)
        assert document.version(version_id).indexing_state is IndexingState.INDEXED
        generation = await pipeline.generations.generation("clearview", version_id)
        assert generation is not None and generation.status is GenerationStatus.COMPLETE

    asyncio.run(scenario())


def test_a_partial_index_write_is_detected_and_never_called_success() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        _, version_id = await pipeline.upload_and_stage()
        await pipeline.make_indexable(version_id)

        class DroppingIndex(InMemorySearchIndex):
            async def index_chunks(self, chunks: Sequence[IndexedChunk]) -> int:
                # Simulates the index accepting only part of the batch.
                return await super().index_chunks(chunks[:-1])

        pipeline.dependencies.index = DroppingIndex()

        with pytest.raises(JobExecutionError) as captured:
            await pipeline.run_job("clearview", version_id)
        assert captured.value.error_code == "ingestion_chunk_write_mismatch"
        assert captured.value.retryable

        document = await pipeline.knowledge.document_for_version("clearview", version_id)
        assert document.version(version_id).indexing_state is IndexingState.FAILED

    asyncio.run(scenario())


def test_a_document_that_fails_scanning_is_a_permanent_refusal() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        _, version_id = await pipeline.upload_and_stage(content=b"# title\n\n\x00binary\x00")
        await pipeline.make_indexable(version_id)

        with pytest.raises(JobExecutionError) as captured:
            await pipeline.run_job("clearview", version_id)
        assert captured.value.error_code == "ingestion_scan_rejected"
        assert not captured.value.retryable
        assert not await pipeline.retrievable()

    asyncio.run(scenario())


def test_an_index_outage_fails_retryably() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        _, version_id = await pipeline.upload_and_stage()
        await pipeline.make_indexable(version_id)

        class UnreachableIndex(InMemorySearchIndex):
            async def index_chunks(self, chunks: Sequence[IndexedChunk]) -> int:
                # Simulates a cluster outage mid-write.
                raise SearchIndexOperationError("search index unreachable")

        pipeline.dependencies.index = UnreachableIndex()

        with pytest.raises(JobExecutionError) as captured:
            await pipeline.run_job("clearview", version_id)
        assert captured.value.error_code == "ingestion_index_unavailable"
        assert captured.value.retryable

    asyncio.run(scenario())


def test_indexing_a_new_version_deactivates_the_superseded_versions_chunks() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        document_id, first_version = await pipeline.upload_and_stage(content=b"2026 terms")
        await pipeline.make_indexable(first_version)
        await pipeline.run_job("clearview", first_version)

        same_document, second_version = await pipeline.upload_and_stage(
            content=b"2027 terms", external_key="financing-options.md"
        )
        assert same_document == document_id
        await pipeline.make_indexable(second_version, at=NOW)
        await pipeline.run_job("clearview", second_version)

        active = await pipeline.index.active_version_ids(
            tenant_id="clearview", document_id=document_id
        )
        assert active == (second_version,)
        assert (
            await pipeline.index.active_chunk_count(tenant_id="clearview", version_id=first_version)
            == 0
        )

    asyncio.run(scenario())


def test_chunking_is_deterministic_and_bounded() -> None:
    text = " ".join(f"token-{index}" for index in range(500))
    chunks = chunk_text(text)
    assert chunks
    # The window budget is measured in estimated characters, so 500 tokens
    # span more than one window — but every window stays within the token
    # budget and repeated runs are identical.
    assert len(chunks) > 1
    assert all(len(chunk.split()) <= 650 for chunk in chunk_text(text * 3))
    assert chunk_text(text) == chunk_text(text)


def test_scanning_rejects_corrupt_and_hostile_content() -> None:
    with pytest.raises(ValidationError):
        parse_document(b"\xff\xfe not utf8", media_type="text/plain", title="T")
    with pytest.raises(ValidationError):
        parse_document(b"# title\n\n\x00null bytes", media_type="text/markdown", title="T")
    with pytest.raises(ValidationError):
        parse_document(b"   \n  ", media_type="text/plain", title="T")
    with pytest.raises(ValidationError):
        parse_document(b"anything", media_type="application/x-msdownload", title="T")


def test_the_parser_is_markdown_section_aware() -> None:
    parsed = parse_document(
        b"# Financing\n\nSome terms.\n\n## Rates\n\n4.9% APR.",
        media_type="text/markdown",
        title="Terms",
    )
    assert parsed.title == "Terms"
    assert parsed.blocks[0].location.section_path == ("Financing",)
    assert parsed.blocks[1].location.section_path == ("Financing", "Rates")
    assert "4.9% APR." in parsed.blocks[1].text


def test_large_documents_are_embedded_in_bounded_batches() -> None:
    """RAG-003 acceptance: batches stay below the per-request embedding limit."""

    async def scenario() -> None:
        pipeline = Pipeline()
        # ~14 000 tokens: far more chunks than a single batch may carry.
        content = (
            b"# Terms\n\n" + (" ".join(f"term-{index}" for index in range(14000)).encode()) + b"\n"
        )
        _, version_id = await pipeline.upload_and_stage(content=content)
        await pipeline.make_indexable(version_id)

        await pipeline.run_job("clearview", version_id)

        assert len(pipeline.embedder.calls) > 1
        assert all(len(call) <= EMBED_BATCH_SIZE for call in pipeline.embedder.calls)
        generation = await pipeline.generations.generation("clearview", version_id)
        assert generation is not None
        assert generation.status is GenerationStatus.COMPLETE
        assert generation.parser_version == "markdown.v1"
        assert generation.chunker_version == "token-window.v2"

    asyncio.run(scenario())


@pytest.mark.parametrize(
    ("content", "media_type", "parser_version"),
    [
        (b"# Heading\n\nPlain paragraph.", "text/markdown", "markdown.v1"),
        (b"<h1>Heading</h1>\n<p>Plain paragraph.</p>", "text/html", "html.v1"),
        (b"Plain paragraph.", "text/plain", "text.v1"),
    ],
)
def test_every_supported_text_format_is_ingested_end_to_end(
    content: bytes, media_type: str, parser_version: str
) -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        _, version_id = await pipeline.upload_and_stage(content=content, media_type=media_type)
        await pipeline.make_indexable(version_id)

        await pipeline.run_job("clearview", version_id)

        generation = await pipeline.generations.generation("clearview", version_id)
        assert generation is not None
        assert generation.status is GenerationStatus.COMPLETE
        assert generation.parser_version == parser_version
        assert generation.chunker_version == "token-window.v2"
        assert await pipeline.retrievable()

    asyncio.run(scenario())


def test_a_pdf_is_ingested_end_to_end_with_page_locations() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        content = _two_page_pdf()
        _, version_id = await pipeline.upload_and_stage(
            content=content, media_type="application/pdf"
        )
        await pipeline.make_indexable(version_id)

        await pipeline.run_job("clearview", version_id)

        generation = await pipeline.generations.generation("clearview", version_id)
        assert generation is not None
        assert generation.status is GenerationStatus.COMPLETE
        assert generation.parser_version == "pdf.v1"
        sections = {chunk.section for chunk in pipeline.index._chunks.values()}
        assert sections == {"Section One (p. 1)", "Section Two (p. 2)"}

    asyncio.run(scenario())


def test_a_docx_is_ingested_end_to_end_with_heading_hierarchy() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        content = _docx_bytes()
        _, version_id = await pipeline.upload_and_stage(
            content=content, media_type=_DOCX_MEDIA_TYPE
        )
        await pipeline.make_indexable(version_id)

        await pipeline.run_job("clearview", version_id)

        generation = await pipeline.generations.generation("clearview", version_id)
        assert generation is not None
        assert generation.status is GenerationStatus.COMPLETE
        assert generation.parser_version == "docx.v1"
        sections = {chunk.section for chunk in pipeline.index._chunks.values()}
        assert sections == {"Terms and conditions", "Terms and conditions > Rates"}

    asyncio.run(scenario())


def test_an_unsupported_media_type_fails_the_job_permanently() -> None:
    async def scenario() -> None:
        pipeline = Pipeline()
        _, version_id = await pipeline.upload_and_stage(
            content=b"MZ\x90\x00", media_type="application/x-msdownload"
        )
        await pipeline.make_indexable(version_id)

        with pytest.raises(JobExecutionError) as captured:
            await pipeline.run_job("clearview", version_id)
        assert captured.value.error_code == "ingestion_scan_rejected"
        assert not captured.value.retryable
        assert not await pipeline.retrievable()

    asyncio.run(scenario())


def _two_page_pdf() -> bytes:
    import io

    from pypdf import PdfWriter
    from pypdf.generic import DictionaryObject, NameObject, StreamObject

    writer = PdfWriter()
    for number, text in enumerate(("Page one text.", "Page two text.")):
        page = writer.add_blank_page(width=612, height=792)
        stream = StreamObject()
        stream.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode())
        page[NameObject("/Contents")] = stream
        font = DictionaryObject(
            {
                NameObject("/F1"): DictionaryObject(
                    {
                        NameObject("/Type"): NameObject("/Font"),
                        NameObject("/Subtype"): NameObject("/Type1"),
                        NameObject("/BaseFont"): NameObject("/Helvetica"),
                    }
                )
            }
        )
        page[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"): font})
        if number == 1:
            writer.add_outline_item("Section Two", 1)
    writer.add_outline_item("Section One", 0)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


_DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx_bytes() -> bytes:
    import io

    from docx import Document

    document = Document()
    document.add_heading("Terms and conditions", level=1)
    document.add_paragraph("First paragraph.")
    document.add_heading("Rates", level=2)
    document.add_paragraph("4.9% APR.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
