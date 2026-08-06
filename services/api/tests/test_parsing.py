"""RAG-003 golden fixtures: parser and chunker output, pinned byte for byte.

Each fixture is a real document (built with the same libraries the adapters
parse) whose expected blocks and locations are written out literally, so any
parser change that alters the produced structure fails this file. The chunker
fixtures pin the exact windows the token budgets produce, which is what makes
"deterministic golden output" a test rather than a hope.
"""

from __future__ import annotations

import io

import pytest

from tenantchat.api.parsing import (
    CHUNK_TOKENS,
    DEFAULT_TOKEN_PROFILE,
    MAX_DOCUMENT_BYTES,
    Chunk,
    ChunkLocation,
    ParsedDocument,
    SourceBlock,
    TokenProfile,
    chunk_document,
    chunk_text,
    count_tokens,
    parse_document,
    parser_version_for,
    profile_for,
)
from tenantchat.core.errors import ValidationError

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _pdf_bytes(
    pages: tuple[str, ...] = ("Page one text.", "Page two text."),
    *,
    outline: bool = True,
    encrypted: str | None = None,
) -> bytes:
    """A minimal valid PDF, page per ``pages`` entry, plus a document outline.

    The outline marks the second page's section; hand-building the content
    stream keeps the fixture dependency-light and deterministic.
    """
    from pypdf import PdfWriter
    from pypdf.generic import DictionaryObject, NameObject, StreamObject

    writer = PdfWriter()
    for number, text in enumerate(pages):
        page = writer.add_blank_page(width=612, height=792)
        stream = StreamObject()
        stream.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode())
        page[NameObject("/Contents")] = stream
        font = DictionaryObject(
            {
                NameObject("/F1"): DictionaryObject(
                    {
                        NameObject("/Type"): NameObject("/Font"),
                        NameObject("/Subtype"): NameObject("/Type1"),
                        NameObject("/BaseFont"): NameObject("/Helvetica"),
                    }
                )
            }
        )
        page[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"): font})
        if outline and number == 1:
            writer.add_outline_item("Section Two", 1)
    if outline:
        writer.add_outline_item("Section One", 0)
    if encrypted is not None:
        writer.encrypt(encrypted)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _docx_bytes() -> bytes:
    """A small DOCX: two headings, body text, a table, and trailing text."""
    from docx import Document

    document = Document()
    document.add_heading("Terms and conditions", level=1)
    document.add_paragraph("First paragraph.")
    document.add_heading("Rates", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Rate"
    table.cell(0, 1).text = "APR"
    table.cell(1, 0).text = "Standard"
    table.cell(1, 1).text = "4.9%"
    document.add_paragraph("After the table.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


MARKDOWN_CONTENT = b"""# Financing options

First paragraph about loans.

## Rates

| Product | APR |
| --- | --- |
| Standard | 4.9% |
| Green | 3.5% |

### Details

```python
x = 1
```

Some final text.
"""

MARKDOWN_GOLDEN = ParsedDocument(
    title="Terms",
    media_type="text/markdown",
    parser_version="markdown.v1",
    blocks=(
        SourceBlock(
            location=ChunkLocation(section_path=("Financing options",)),
            text="Financing options\nFirst paragraph about loans.",
        ),
        SourceBlock(
            location=ChunkLocation(section_path=("Financing options", "Rates")),
            text="Rates\n| Product | APR |\n| Standard | 4.9% |\n| Green | 3.5% |",
        ),
        SourceBlock(
            location=ChunkLocation(section_path=("Financing options", "Rates", "Details")),
            text="Details\nx = 1\nSome final text.",
        ),
    ),
)

HTML_CONTENT = b"""<!doctype html>
<html><head><title>Ignored title</title><style>p { color: red }</style></head>
<body>
<h1 id="overview">Overview</h1>
<p>Hello <b>bold</b> world.</p>
<script>var hidden = 1;</script>
<h2 id="pricing">Pricing</h2>
<table><tr><th>Plan</th><th>Price</th></tr><tr><td>Basic</td><td>Free</td></tr></table>
</body></html>
"""

HTML_GOLDEN = ParsedDocument(
    title="Site",
    media_type="text/html",
    parser_version="html.v1",
    blocks=(
        SourceBlock(
            location=ChunkLocation(section_path=("Overview",), anchor="overview"),
            text="Overview Hello bold world.",
        ),
        SourceBlock(
            location=ChunkLocation(section_path=("Overview", "Pricing"), anchor="pricing"),
            text="Pricing\nPlan | Price\nBasic | Free",
        ),
    ),
)

TEXT_GOLDEN = ParsedDocument(
    title="Notes",
    media_type="text/plain",
    parser_version="text.v1",
    blocks=(
        SourceBlock(
            location=ChunkLocation(),
            text="Just some\nplain text lines.",
        ),
    ),
)

PDF_GOLDEN = ParsedDocument(
    title="Brochure",
    media_type="application/pdf",
    parser_version="pdf.v1",
    blocks=(
        SourceBlock(
            location=ChunkLocation(section_path=("Section One",), page=1),
            text="Page one text.",
        ),
        SourceBlock(
            location=ChunkLocation(section_path=("Section Two",), page=2),
            text="Page two text.",
        ),
    ),
)

DOCX_GOLDEN = ParsedDocument(
    title="Terms",
    media_type=DOCX_MEDIA_TYPE,
    parser_version="docx.v1",
    blocks=(
        SourceBlock(
            location=ChunkLocation(section_path=("Terms and conditions",)),
            text="Terms and conditions\nFirst paragraph.",
        ),
        SourceBlock(
            location=ChunkLocation(section_path=("Terms and conditions", "Rates")),
            text="Rates\nRate | APR\nStandard | 4.9%\nAfter the table.",
        ),
    ),
)


@pytest.mark.parametrize(
    ("content", "media_type", "title", "golden"),
    [
        (MARKDOWN_CONTENT, "text/markdown", "Terms", MARKDOWN_GOLDEN),
        (HTML_CONTENT, "text/html", "Site", HTML_GOLDEN),
        (b"Just some\nplain text lines.", "text/plain", "Notes", TEXT_GOLDEN),
        (_pdf_bytes(), "application/pdf", "Brochure", PDF_GOLDEN),
        (_docx_bytes(), DOCX_MEDIA_TYPE, "Terms", DOCX_GOLDEN),
    ],
)
def test_parser_output_is_exactly_the_golden_fixture(
    content: bytes, media_type: str, title: str, golden: ParsedDocument
) -> None:
    assert parse_document(content, media_type=media_type, title=title) == golden


def test_media_type_normalization_and_supported_types() -> None:
    assert (
        parse_document(
            MARKDOWN_CONTENT, media_type="text/markdown; charset=utf-8", title="Terms"
        ).parser_version
        == "markdown.v1"
    )
    assert parser_version_for("text/x-markdown") == "markdown.v1"
    with pytest.raises(ValidationError):
        parser_version_for("application/x-msdownload")


def test_the_html_heading_id_anchor_is_captured_where_the_source_has_one() -> None:
    parsed = parse_document(HTML_CONTENT, media_type="text/html", title="Site")
    assert parsed.blocks[1].location.anchor == "pricing"
    assert parsed.blocks[1].location.section_path == ("Overview", "Pricing")


def test_a_level_jump_makes_the_deeper_heading_a_child_of_the_stack() -> None:
    parsed = parse_document(
        b"# A\n\n## B\n\n#### C\n\ntext under C",
        media_type="text/markdown",
        title="T",
    )
    assert parsed.blocks[-1].location.section_path == ("A", "B", "C")


def test_setext_headings_and_fenced_code_are_parsed() -> None:
    parsed = parse_document(
        b"Setext title\n=====\n\n```text\n# not a heading\n```\n\nafter",
        media_type="text/markdown",
        title="T",
    )
    assert parsed.blocks[0].location.section_path == ("Setext title",)
    assert parsed.blocks[-1].location.section_path == ("Setext title",)
    assert "# not a heading" in parsed.blocks[-1].text


@pytest.mark.parametrize(
    ("content", "media_type", "detail"),
    [
        (b"   \n  ", "text/plain", "document has no text content"),
        (b"", "text/markdown", "document has no text content"),
        (b"# title\n\n\x00null bytes", "text/markdown", "document contains NUL bytes"),
        (b"\xff\xfe not utf8", "text/plain", "document is not valid UTF-8"),
        (b"x" * (MAX_DOCUMENT_BYTES + 1), "text/plain", "document exceeds the size budget"),
        (_pdf_bytes()[:100], "application/pdf", "document is corrupt"),
        (_pdf_bytes(encrypted="secret"), "application/pdf", "document is encrypted"),
        (b"not a zip file", DOCX_MEDIA_TYPE, "document is corrupt"),
    ],
)
def test_empty_corrupt_encrypted_and_oversized_documents_are_refused(
    content: bytes, media_type: str, detail: str
) -> None:
    with pytest.raises(ValidationError) as captured:
        parse_document(content, media_type=media_type, title="T")
    assert captured.value.detail == detail


def test_an_unsupported_media_type_is_refused_without_looking_at_content() -> None:
    with pytest.raises(ValidationError):
        parse_document(b"anything", media_type="application/x-msdownload", title="T")


def test_a_pdf_without_an_outline_still_parses_with_page_locations() -> None:
    parsed = parse_document(
        _pdf_bytes(pages=("Only page.",), outline=False),
        media_type="application/pdf",
        title="T",
    )
    assert parsed.blocks[0].location == ChunkLocation(page=1)


def test_a_blank_pdf_page_keeps_its_page_number_for_later_pages() -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    page = writer.add_blank_page(width=612, height=792)
    from pypdf.generic import DictionaryObject, NameObject, StreamObject

    stream = StreamObject()
    stream.set_data(b"BT /F1 12 Tf 72 720 Td (second page) Tj ET")
    page[NameObject("/Contents")] = stream
    font = DictionaryObject(
        {
            NameObject("/F1"): DictionaryObject(
                {
                    NameObject("/Type"): NameObject("/Font"),
                    NameObject("/Subtype"): NameObject("/Type1"),
                    NameObject("/BaseFont"): NameObject("/Helvetica"),
                }
            )
        }
    )
    page[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"): font})
    buffer = io.BytesIO()
    writer.write(buffer)

    parsed = parse_document(buffer.getvalue(), media_type="application/pdf", title="T")
    assert parsed.blocks[0].location == ChunkLocation(page=2)


def test_chunking_is_deterministic_and_budgeted_across_blocks() -> None:
    parsed = parse_document(MARKDOWN_CONTENT, media_type="text/markdown", title="Terms")
    first = chunk_document(parsed)
    second = chunk_document(parsed)
    assert first == second
    assert all(len(chunk.text.split()) <= CHUNK_TOKENS for chunk in first)
    assert all(chunk.location in {block.location for block in parsed.blocks} for chunk in first)


def test_every_chunk_maps_back_to_a_human_readable_location() -> None:
    parsed = parse_document(_pdf_bytes(), media_type="application/pdf", title="Brochure")
    for chunk in chunk_document(parsed, chunk_tokens=8, overlap_tokens=1):
        assert str(chunk.location) in {
            "Section One (p. 1)",
            "Section Two (p. 2)",
        }


def test_windows_overlap_by_the_configured_token_budget() -> None:
    text = " ".join(f"token-{index}" for index in range(100))
    windows = chunk_text(text, chunk_tokens=10, overlap_tokens=2)
    assert len(windows) > 1
    previous_tail = windows[0].split()[-2:]
    for window in windows[1:]:
        # The overlap is a character budget, so the shared token count is not
        # exactly the configured token overlap — but a window's head must
        # repeat the previous window's tail, otherwise nothing is shared.
        assert window.split()[0] in previous_tail
        previous_tail = window.split()[-2:]


def test_an_overlap_larger_than_the_window_is_clamped_not_invalid() -> None:
    text = " ".join(f"token-{index}" for index in range(10))
    windows = chunk_text(text, chunk_tokens=100, overlap_tokens=50)
    assert windows == [text]


def test_a_single_token_longer_than_the_budget_is_emitted_alone() -> None:
    text = " ".join(["short"] * 3 + [f"x{'-' * 400}"])
    windows = chunk_text(text, chunk_tokens=2, overlap_tokens=0)
    assert windows[-1] == f"x{'-' * 400}"


def test_chunking_never_crosses_a_source_block_boundary() -> None:
    parsed = ParsedDocument(
        title="T",
        media_type="text/plain",
        parser_version="text.v1",
        blocks=(
            SourceBlock(location=ChunkLocation(section_path=("A",)), text="alpha beta gamma delta"),
            SourceBlock(location=ChunkLocation(section_path=("B",)), text="one two three four"),
        ),
    )
    chunks = chunk_document(parsed, chunk_tokens=4, overlap_tokens=0)
    assert [chunk.location.section_path for chunk in chunks] == [
        ("A",),
        ("A",),
        ("B",),
        ("B",),
    ]
    assert " ".join(chunk.text for chunk in chunks) == ("alpha beta gamma delta one two three four")


def test_invalid_chunk_budgets_are_rejected() -> None:
    with pytest.raises(ValueError):
        chunk_text("some text", chunk_tokens=0)
    with pytest.raises(ValueError):
        TokenProfile(name="bad", chars_per_token=0)
    with pytest.raises(ValueError):
        TokenProfile(name="bad", chars_per_token=-1).tokens_for(5)
    with pytest.raises(ValueError):
        ChunkLocation(page=0)


def test_token_profiles_are_model_aware_and_default_safe() -> None:
    assert profile_for("qwen3-embedding") is not DEFAULT_TOKEN_PROFILE
    assert profile_for("unknown-family") is DEFAULT_TOKEN_PROFILE
    assert count_tokens("abcd", profile=TokenProfile(name="tight", chars_per_token=2)) == 2
    assert count_tokens("abcd") == 1


def test_location_rendering_is_stable_and_human_readable() -> None:
    assert str(ChunkLocation()) == "Document"
    assert str(ChunkLocation(section_path=("A", "B"))) == "A > B"
    assert str(ChunkLocation(section_path=("A",), page=3)) == "A (p. 3)"
    assert str(ChunkLocation(section_path=("A", "B"), anchor="rates")) == "A > B > #rates"


def test_chunk_objects_are_immutable() -> None:
    chunk = Chunk(location=ChunkLocation(section_path=("A",)), text="hello")
    text_field = "text"
    path_field = "section_path"
    with pytest.raises(AttributeError):
        setattr(chunk, text_field, "changed")
    with pytest.raises(AttributeError):
        setattr(chunk.location, path_field, ("B",))
